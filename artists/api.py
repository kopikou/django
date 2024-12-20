from rest_framework.viewsets import GenericViewSet
from rest_framework import mixins, viewsets

from artists import serializers
from artists.models import Artist,Show,Type,Income,Expenses
from artists.serializers import *
from django.contrib.auth.models import User
from rest_framework.decorators import action
from rest_framework.response import Response
from django.db.models import Avg, Count, Max, Min 
from django.contrib.auth import authenticate, login, logout
from django.core.cache import cache
from rest_framework.permissions import IsAuthenticated, BasePermission
from rest_framework.views import APIView
from rest_framework import status
import pandas as pd
from django.http import HttpResponse;
from docx import Document; 

from openpyxl import Workbook
from openpyxl.utils import get_column_letter
import io


class UsersViewset(mixins.RetrieveModelMixin, 
    mixins.ListModelMixin, 
    GenericViewSet):
    queryset = User.objects.all()
    serializer_class = UserSerializer

class UserProfileViewSet(GenericViewSet):
    permission_classes = [IsAuthenticated]

    @action(detail=False, methods=['get'], url_path='info')
    def get_info(self, request, *args, **kwargs):
        user = request.user
        data = {
            "is_authenticated": user.is_authenticated,
        }
        if user.is_authenticated:
            data.update({
                "is_superuser": user.is_superuser,
                "id": user.id
            })
        return Response(data)
    
    @action(detail=False, methods=['post'], url_path='login', permission_classes=[])
    def login(self, request, *args, **kwargs):
        username = self.request.data['username']
        password = self.request.data['password']

        user = authenticate(username=username, password=password)
        if user:
            login(request, user)
        
        return Response({
            'is_auth': bool(user)
        })
    
    @action(detail=False, methods=['get'], url_path='logout', permission_classes=[])
    def logout(self, request, *args, **kwargs):
        # username = self.request.data['username']
        # password = self.request.data['password']

        # user = authenticate(username=username, password=password)
        # if user:
        #     login(request, user)
        
        # return Response({
        #     'is_auth': bool(user)
        # })
        if self.request.user.is_authenticated:
            #request.user.auth_token.delete()
            logout(request)
        return Response({
            'is_auth': False
        })
    

class ArtistsViewset(mixins.CreateModelMixin,
    mixins.UpdateModelMixin, 
    mixins.RetrieveModelMixin,                
    mixins.ListModelMixin,
    mixins.DestroyModelMixin,
    GenericViewSet):
    queryset = Artist.objects.all()
    serializer_class = ArtistSerializer

    def get_serializer_class(self):
        if self.action in ('create', 'update'):
            return ArtistCreateSerializer
        return super().get_serializer_class()
    def get_queryset(self):
        qs = super().get_queryset()

        show = self.request.query_params.get('show')
        name = self.request.query_params.get('name')
        #user = self.request.query_params.get('user')
        if show:
            qs = qs.filter(show=show)
        if name:
            qs = qs.filter(name=name)
        # if user:
        #     qs = qs.filter(user=user)
        # if not self.request.user.is_superuser:
        #     qs = qs.filter(user=self.request.user)
        if(not self.request.user.is_authenticated):
            #qs = qs.filter(user=self.request.user)
            qs = None
        return qs

class ShowViewset(mixins.CreateModelMixin,
    mixins.UpdateModelMixin, 
    mixins.RetrieveModelMixin,                
    mixins.ListModelMixin,
    mixins.DestroyModelMixin,
    GenericViewSet):
    queryset = Show.objects.all()
    serializer_class = ShowSerializer

    def get_serializer_class(self):
        if self.action in ('create', 'update'):
            return ShowCreateSerializer
        return super().get_serializer_class()
    def get_queryset(self):
        qs = super().get_queryset()

        type = self.request.query_params.get('type')

        if type:
            qs = qs.filter(type=type)
        if(not self.request.user.is_authenticated):
            #qs = qs.filter(user=self.request.user)
            qs = None
        return qs
    class StatsSerializer(serializers.Serializer):
        count = serializers.IntegerField()
        avg = serializers.FloatField()
        max = serializers.IntegerField()
        min = serializers.IntegerField()

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, request, *args, ** kwargs):
        stats = Show.objects.aggregate(
        count=Count("*"),
        avg=Avg("price"),
        max=Max("price"),
        min=Min("price"),
        )
        serializer = self.StatsSerializer(instance=stats)
        return Response(serializer.data)

class TypeViewset(mixins.CreateModelMixin,
    mixins.UpdateModelMixin, 
    mixins.RetrieveModelMixin,                
    mixins.ListModelMixin,
    mixins.DestroyModelMixin,
    GenericViewSet):
    queryset = Type.objects.all()
    serializer_class = TypeSerializer
    def get_queryset(self):
        qs = super().get_queryset()

        if(not self.request.user.is_authenticated):
            #qs = qs.filter(user=self.request.user)
            qs = None
        return qs

class IncomeViewset(mixins.CreateModelMixin,
    mixins.UpdateModelMixin, 
    mixins.RetrieveModelMixin,                
    mixins.ListModelMixin,
    mixins.DestroyModelMixin,
    GenericViewSet):
    queryset = Income.objects.all()
    serializer_class = IncomeSerializer

    def get_serializer_class(self):
        if self.action in ('create', 'update'):
            return IncomeCreateSerializer
        return super().get_serializer_class()
    
    def get_queryset(self):
        qs = super().get_queryset()

        show = self.request.query_params.get('show')
        user = self.request.query_params.get('user')
        if show:
            qs = qs.filter(show=show)
        if user:
            qs = qs.filter(user=user)
        if not self.request.user.is_superuser and self.request.user.is_authenticated:
            qs = qs.filter(user=self.request.user)
        if(not self.request.user.is_authenticated):
            qs = None
        return qs
    
    
class ExpenseViewset(mixins.CreateModelMixin,
    mixins.UpdateModelMixin, 
    mixins.RetrieveModelMixin,                
    mixins.ListModelMixin,
    mixins.DestroyModelMixin,
    GenericViewSet):
    queryset = Expenses.objects.all()
    serializer_class = ExpensesSerializer    

    def get_serializer_class(self):
        if self.action in ('create', 'update'):
            return ExpensesCreateSerializer
        return super().get_serializer_class()
    def get_queryset(self):
        qs = super().get_queryset()

        artist = self.request.query_params.get('artist')
        user = self.request.query_params.get('user')
        income = self.request.query_params.get('income')
        if artist:
            qs = qs.filter(artist=artist)
        if income:
            qs = qs.filter(income=income)
        if user:
            qs = qs.filter(user=user)
        if not self.request.user.is_superuser and self.request.user.is_authenticated:
            qs = qs.filter(user=self.request.user)
        if(not self.request.user.is_authenticated):
            qs = None
        return qs
    class StatsSerializer(serializers.Serializer):
        count = serializers.IntegerField()
        avg = serializers.FloatField()
        max = serializers.IntegerField()
        min = serializers.IntegerField()

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, request, *args, ** kwargs):
        stats = Expenses.objects.aggregate(
        count=Count("*"),
        avg=Avg("salary"),
        max=Max("salary"),
        min=Min("salary"),
        )
        serializer = self.StatsSerializer(instance=stats)
        return Response(serializer.data)
    
    class ExportExpensesSerializer(serializers.ModelSerializer):
        artist_name = serializers.CharField(source='artist.name')  
        income_id = serializers.IntegerField(source='income.id')
        show = serializers.CharField(source='income.show')
        date = serializers.DateField(source='income.date')
        salary_pers = serializers.IntegerField(source='salary')

        class Meta:
            model = Expenses
            fields = ['id', 'artist_name', 'income_id', 'show', 'date', 'salary_pers']

    @action(detail=False, methods=['GET'], url_path='export-excel')
    def export_excel(self, request):
        queryset = self.get_queryset()
        artist = self.request.query_params.get('artist')
        user = self.request.query_params.get('user')
        income = self.request.query_params.get('income')
        if artist:
            queryset = queryset.filter(artist=artist)
        if income:
            queryset = queryset.filter(income=income)
        if user:
            queryset = queryset.filter(user=user)
        if not self.request.user.is_superuser and self.request.user.is_authenticated:
            queryset = queryset.filter(user=self.request.user)
        if(not self.request.user.is_authenticated):
            queryset = None

        serializer = self.ExportExpensesSerializer(queryset, many=True)
        data = serializer.data

        df = pd.DataFrame(data)

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = 'attachment; filename="Расходы.xlsx"'
        df.to_excel(response, index=False)
        return response
    @action(detail=False, methods=['GET'], url_path='export-word')
    def export_word(self, request):
        queryset = self.get_queryset()
        artist = self.request.query_params.get('artist')
        user = self.request.query_params.get('user')
        income = self.request.query_params.get('income')
        if artist:
            queryset = queryset.filter(artist=artist)
        if income:
            queryset = queryset.filter(income=income)
        if user:
            queryset = queryset.filter(user=user)
        if not self.request.user.is_superuser and self.request.user.is_authenticated:
            queryset = queryset.filter(user=self.request.user)
        if(not self.request.user.is_authenticated):
            queryset = None

        serializer = self.ExportExpensesSerializer(queryset, many=True)
        data = serializer.data

        doc = Document()
        doc.add_heading('Данные о расходах', level=1)

        for item in data:
            doc.add_paragraph(str(item))

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Расходы.docx"'
        doc.save(response)
        return response
